from docx import Document

def ler_docx(file_path):
    doc = Document(file_path)
    paragrafos = []

    # for section in doc.sections:
    #     for paragraph in section.header.paragraphs:
    #         paginas.append(paragraph.text)
    #     for paragraph in section.footer.paragraphs:
    #         paginas.append(paragraph.text)

    i = 0
    paginas = []
    for paragraph in doc.paragraphs:
        paragrafos.append(paragraph.text+'\n')
        if i >= 50:
            paginas.append(''.join(paragrafos))
            paragrafos.clear()
            i = 0
        else:
            i += 1
    
    if i > 0:
        paginas.append(''.join(paragrafos))

    return paginas